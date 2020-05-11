#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# load tempfile for temporary dir creation
import sys, os, tempfile

import json

# load libraries for NLP pipeline
import spacy
# load Visualizers 
from spacy import displacy

# load misc utils
import uuid
from werkzeug.utils import secure_filename
import logging
# logging.basicConfig(format='%(asctime)s : %(levelname)s : %(message)s', level=logging.INFO)
logging.basicConfig(format='%(asctime)s : %(levelname)s : %(message)s', level=logging.DEBUG)
# logging.basicConfig(format='%(asctime)s : %(levelname)s : %(message)s', level=logging.ERROR)

# load libraries for string proccessing
import re, string

# load libraries for API proccessing
from flask import Flask, jsonify, flash, request, Response, redirect, url_for, abort, render_template

# A Flask extension for handling Cross Origin Resource Sharing (CORS), making cross-origin AJAX possible.
from flask_cors import CORS

# python-docx
# https://python-docx.readthedocs.io/en/latest/index.html
from docx import Document

# load libraries for XML proccessing
import xml.etree.ElementTree as ET

ALLOWED_EXTENSIONS = set(['docx']) 

# Load globally spaCy model via package name
NLP_NB = spacy.load('nb_core_news_sm')
NLP_NB_VECTORES = spacy.load('./tmp/nb_nowac_vectores')
# NLP_EN_VECTORES = spacy.load('en_core_web_lg')

# load SnowballStemmer stemmer from nltk
from nltk.stem.snowball import SnowballStemmer
# Load globally english SnowballStemmer
NORWEGIAN_STEMMER = SnowballStemmer("norwegian")

__author__ = "Kyrylo Malakhov <malakhovks@nas.gov.ua> and Vitalii Velychko <aduisukr@gmail.com>"
__copyright__ = "Copyright (C) 2020 Kyrylo Malakhov <malakhovks@nas.gov.ua> and Vitalii Velychko <aduisukr@gmail.com>"

# class JSONResponse(Response):
#     default_mimetype = 'application/json'

app = Flask(__name__)
CORS(app)
# app.response_class = JSONResponse

"""
Limited the maximum allowed payload to 16 megabytes.
If a larger file is transmitted, Flask will raise an RequestEntityTooLarge exception.
"""
# app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

"""
Set the secret key to some random bytes. Keep this really secret!
How to generate good secret keys.
A secret key should be as random as possible. Your operating system has ways to generate pretty random data based on a cryptographic random generator. Use the following command to quickly generate a value for Flask.secret_key (or SECRET_KEY):
$ python -c 'import os; print(os.urandom(16))'
b'_5#y2L"F4Q8z\n\xec]/'
"""
# app.secret_key = b'_5#y2L"F4Q8z\n\xec]/'
app.secret_key = os.urandom(42)

"""
# ------------------------------------------------------------------------------------------------------
# Functions
# ------------------------------------------------------------------------------------------------------
# """

# function that check if an extension is valid
def allowed_file(filename):
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Ridge DOCX-templated CV processing into JSON-----------------------------------------------------------
def get_data_from_ridge_docx_into_json(docx_path):
    """
    Take the path of a docx file as argument, return the text in unicode.
    """
    cv_data = {}
    personal_name_cv = {}
    paragraphs_temp = []
    document = Document(docx_path)

    # Personal name extraction
    personal_name_cv['First name'] = document.core_properties.subject.split(' ', 1)[0]
    personal_name_cv['Last name'] = document.core_properties.subject.split(personal_name_cv['First name'] + ' ', 1)[1]
    cv_data['Personal name'] = personal_name_cv

    # First table extraction BIO
    table = document.tables[0]
    bio_data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        # Construct a tuple for this row
        # if contain letters and numbers
        row_data = tuple(x for x in text if re.search(r'[A-Za-z0-9]+', x))
        bio_data.append(row_data)

    # remove an empty tuple(s) from a list of tuples
    bio_data = [t for t in bio_data if t]

    logging.debug(bio_data)

    cv_data['Bio'] = bio_data

    # Key Skills extraction
    paragraphs_temp.append(''.join(document.core_properties.subject))
    for para in document.paragraphs:
        if re.search(r'[A-Za-z0-9]+', para.text):
            text = para.text
            paragraphs_temp.append(''.join(text))
    if 'Work Experience' in paragraphs_temp:
        paragraphs_temp = paragraphs_temp[paragraphs_temp.index('Key Skills')+1:paragraphs_temp.index('Work Experience')]
    if 'Qualifications' in paragraphs_temp:
        paragraphs_temp = paragraphs_temp[paragraphs_temp.index('Key Skills')+1:paragraphs_temp.index('Qualifications')]
    cv_data['Key Skills'] = paragraphs_temp
    json_cv_data = json.dumps(cv_data)

    # return '\n'.join(paragraphs_temp)
    return json_cv_data

# --------------------------------------------------------------------------------------------------------

# Ridge-templated CV (DOCX) processing into XML-------------------------------------------------------------------
def docx_ridge_to_xml(docx_path, flg):
    # Init docx document
    document = Document(docx_path)
    paragraphs_temp = []

    # create root element <Graph>
    root_element_Graph = ET.Element("Graph")
    root_element_Graph.set("nodesize", "10")
    root_element_Graph.set("vspacing", "3")
    root_element_Graph.set("hspacing", "20")
    root_element_Graph.set("padding", "4")
    root_element_Graph.set("guid", "")
    root_element_Graph.set("searchurl", "https://www.google.com.ua/#output=search%26q=###SEARCH###")

    # create <Linkgroups> element
    element_Linkgroups = ET.Element("Linkgroups")
    ET.SubElement(element_Linkgroups, "Group", {"name" : "Default", "color" : "10066329"})

    # create <datagroups> element
    element_datagroups = ET.Element("datagroups")
    # create <datagroup> element as subelements of <datagroups> element
    ET.SubElement(element_datagroups, "datagroup").text = "First Name"
    ET.SubElement(element_datagroups, "datagroup").text = "Last Name"
    ET.SubElement(element_datagroups, "datagroup").text = "Birthday"
    ET.SubElement(element_datagroups, "datagroup").text = "Job Title"
    ET.SubElement(element_datagroups, "datagroup").text = "Location"
    ET.SubElement(element_datagroups, "datagroup").text = "Languages"
    ET.SubElement(element_datagroups, "datagroup").text = "Language"
    ET.SubElement(element_datagroups, "datagroup").text = "Nationality"
    ET.SubElement(element_datagroups, "datagroup").text = "Speak Norwegian"
    ET.SubElement(element_datagroups, "datagroup").text = "Skills"
    ET.SubElement(element_datagroups, "datagroup").text = "Skill"
    ET.SubElement(element_datagroups, "datagroup").text = "Consultancy Firm"

    # create <Edges> element
    element_Edges = ET.Element("Edges")

    # add default edge
    ET.SubElement(element_Edges, "Edge", {"guid" : str(uuid.uuid4()).rpartition('-')[-1], "edgeName" : "", "node1" : document.core_properties.subject, "node2" : "Stuff", "group" : "Default", "istwoway" : "false"})

    # create <Nodes> element
    element_Nodes = ET.Element("Nodes")

    # add default node
    ET.SubElement(element_Nodes, "Node", {"guid" : str(uuid.uuid4()).rpartition('-')[-1], "nodeName" : "Stuff", "nclass" : "", "shape" : "circle", "color" : "13421772", "xPos" : "", "yPos" : "", "font" : "Verdana", "fontsize" : "14"})

    element_Node = ET.Element("Node")
    element_Node.set("guid", str(uuid.uuid4()).rpartition('-')[-1])
    element_Node.set("nodeName", document.core_properties.subject)
    element_Node.set("nclass", "")
    element_Node.set("shape", "square")
    element_Node.set("color", "13421772")
    element_Node.set("xPos", "")
    element_Node.set("yPos", "")
    element_Node.set("Font", "Verdana")
    element_Node.set("fontsize", "14")

    element_Data_Skills = ET.Element("data")
    element_Data_Skills.set("tclass", "Skills")
    element_Data_Skills.set("type", "text")
    element_Data_Skills.set("link", "")

    element_Data_Languages = ET.Element("data")
    element_Data_Languages.set("tclass", "Languages")
    element_Data_Languages.set("type", "text")
    element_Data_Languages.set("link", "")

    # add default node with Ridge value
    ET.SubElement(element_Node, "data", {"tclass": "Consultancy Firm", "type" : "text", "link" : ""}).text = "Ridge"

    # add node with vacancy value
    ET.SubElement(element_Node, "data", {"tclass": "First Name", "type" : "text", "link" : ""}).text = document.core_properties.subject.split(' ', 1)[0]

    ET.SubElement(element_Node, "data", {"tclass": "Last Name", "type" : "text", "link" : ""}).text = document.core_properties.subject.split(document.core_properties.subject.split(' ', 1)[0] + ' ', 1)[1]

    # First table extraction BIO
    table = document.tables[0]
    bio_data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        # Construct a tuple for this row
        # if contain letters and numbers
        row_data = tuple(x for x in text if re.search(r'[A-Za-z0-9]+', x))
        bio_data.append(row_data)
    # remove an empty tuple(s) from a list of tuples
    bio_data = [t for t in bio_data if t]

    for i in bio_data:
        if "Title" in i[0]:
            ET.SubElement(element_Node, "data", {"tclass": "Job Title", "type" : "text", "link" : ""}).text =i[1].strip()
            ET.SubElement(element_Nodes, "Node", {"guid" : str(uuid.uuid4()).rpartition('-')[-1], "nodeName" : i[1], "nclass" : "", "shape" : "circle", "color" : "13421772", "xPos" : "", "yPos" : "", "font" : "Verdana", "fontsize" : "14"})
            ET.SubElement(element_Edges, "Edge", {"guid" : str(uuid.uuid4()).rpartition('-')[-1], "edgeName" : "", "node1" : document.core_properties.subject, "node2" : i[1], "group" : "Default", "istwoway" : "false"})
        if i[0] == "Nationality:":
            ET.SubElement(element_Node, "data", {"tclass": "Nationality", "type" : "text", "link" : ""}).text = i[1]
        if i[0] == "Languages:":
            # ET.SubElement(element_Node, "data", {"tclass": "Languages", "type" : "text", "link" : ""}).text = i[1]
            element_Data_Languages.text = i[1]
            for lang in i[1].split(','):
                ET.SubElement(element_Data_Languages, "data", {"tclass": "Language", "type" : "text", "link" : ""}).text = lang.strip()
            element_Node.append(element_Data_Languages)
        if i[0] == "Date of Birth:":
            ET.SubElement(element_Node, "data", {"tclass": "Birthday", "type" : "text", "link" : ""}).text = i[1]
        if i[0] == "Location:":
            ET.SubElement(element_Node, "data", {"tclass": "Location", "type" : "text", "link" : ""}).text = i[1]
        if '\n' in i[0]:
            list_bio_temp_left = i[0].split("\n", 1)
            list_bio_temp_right = i[1].split("\n", 1)
            for j in list_bio_temp_left:
                if 'Date of Birth' in j:
                    ET.SubElement(element_Node, "data", {"tclass": "Birthday", "type" : "text", "link" : ""}).text = list_bio_temp_right[list_bio_temp_left.index(j)]
                if 'Year of Birth' in j:
                    ET.SubElement(element_Node, "data", {"tclass": "Birthday", "type" : "text", "link" : ""}).text = list_bio_temp_right[list_bio_temp_left.index(j)]
                if 'Location' in j:
                    ET.SubElement(element_Node, "data", {"tclass": "Location", "type" : "text", "link" : ""}).text = list_bio_temp_right[list_bio_temp_left.index(j)]
                if 'Nationality' in j:
                    ET.SubElement(element_Node, "data", {"tclass": "Nationality", "type" : "text", "link" : ""}).text = list_bio_temp_right[list_bio_temp_left.index(j)]
                if 'Languages' in j:
                    # ET.SubElement(element_Node, "data", {"tclass": "Languages", "type" : "text", "link" : ""}).text = list_bio_temp_right[list_bio_temp_left.index(j)]
                    element_Data_Languages.text = list_bio_temp_right[list_bio_temp_left.index(j)]
                    for lang in list_bio_temp_right[list_bio_temp_left.index(j)].split(','):
                        ET.SubElement(element_Data_Languages, "data", {"tclass": "Language", "type" : "text", "link" : ""}).text = lang.strip()
                    element_Node.append(element_Data_Languages)

    # Key Skills extraction
    for para in document.paragraphs:
        if re.search(r'[A-Za-z0-9]+', para.text):
            text = para.text
            paragraphs_temp.append(''.join(text))
    if 'Work Experience' in paragraphs_temp:
        if 'Key Strengths' in paragraphs_temp:
            paragraphs_temp = paragraphs_temp[paragraphs_temp.index('Key Strengths')+1 : paragraphs_temp.index('Work Experience')]
            # ET.SubElement(element_Node, "data", {"tclass": "Skills", "type" : "text", "link" : ""}).text = re.sub('\s\s+', ' ', ' '.join(paragraphs_temp))
            element_Data_Skills.text = re.sub('\s\s+', ' ', ' '.join(paragraphs_temp))
            for temp in paragraphs_temp:
                ET.SubElement(element_Data_Skills, "data", {"tclass": "Skill", "type" : "text", "link" : ""}).text = temp
            element_Node.append(element_Data_Skills)
        if 'Key Skills' in paragraphs_temp:
            paragraphs_temp = paragraphs_temp[paragraphs_temp.index('Key Skills')+1 : paragraphs_temp.index('Work Experience')]
            # ET.SubElement(element_Node, "data", {"tclass": "Skills", "type" : "text", "link" : ""}).text = re.sub('\s\s+', ' ', ' '.join(paragraphs_temp))
            element_Data_Skills.text = re.sub('\s\s+', ' ', ' '.join(paragraphs_temp))
            for temp in paragraphs_temp:
                ET.SubElement(element_Data_Skills, "data", {"tclass": "Skill", "type" : "text", "link" : ""}).text = temp
            element_Node.append(element_Data_Skills)
    if 'Qualifications' in paragraphs_temp:
        if 'Key Strengths' in paragraphs_temp:
            paragraphs_temp = paragraphs_temp[paragraphs_temp.index('Key Strengths')+1 : paragraphs_temp.index('Qualifications')]
            # ET.SubElement(element_Node, "data", {"tclass": "Skills", "type" : "text", "link" : ""}).text = re.sub('\s\s+', ' ', ' '.join(paragraphs_temp))
            element_Data_Skills.text = re.sub('\s\s+', ' ', ' '.join(paragraphs_temp))
            for temp in paragraphs_temp:
                ET.SubElement(element_Data_Skills, "data", {"tclass": "Skill", "type" : "text", "link" : ""}).text = temp
            element_Node.append(element_Data_Skills)
        if 'Key Skills' in paragraphs_temp:
            paragraphs_temp = paragraphs_temp[paragraphs_temp.index('Key Skills')+1 : paragraphs_temp.index('Qualifications')]
            # ET.SubElement(element_Node, "data", {"tclass": "Skills", "type" : "text", "link" : ""}).text = re.sub('\s\s+', ' ', ' '.join(paragraphs_temp))
            element_Data_Skills.text = re.sub('\s\s+', ' ', ' '.join(paragraphs_temp))
            for temp in paragraphs_temp:
                ET.SubElement(element_Data_Skills, "data", {"tclass": "Skill", "type" : "text", "link" : ""}).text = temp
            element_Node.append(element_Data_Skills)

    element_Nodes.append(element_Node)
    root_element_Graph.append(element_Nodes)
    root_element_Graph.append(element_Edges)
    root_element_Graph.append(element_Linkgroups)
    root_element_Graph.append(element_datagroups)

    if flg:
        return ET.tostring(root_element_Graph, encoding='utf8', method='xml')
    else:
        return root_element_Graph

# --------------------------------------------------------------------------------------------------------

# Free-templated CV (DOCX) processing --------------------------------------------------------------------
def get_data_from_free_docx(docx_path):
    """
    Take the path of a docx file as argument, return the text in unicode.
    """
    paragraphs = []
    document = Document(docx_path)
    logging.debug(document.core_properties.subject)

    paragraphs.append(''.join(document.core_properties.subject))
    for para in document.paragraphs:
        text = para.text
        paragraphs.append(''.join(text))
    return '\n\n'.join(paragraphs)

"""
# ------------------------------------------------------------------------------------------------------
# """

@app.route('/')
def index():
    return Response(render_template('index.html'), mimetype='text/html')

@app.route('/ner')
def ner():
    return Response(render_template('ner.html'), mimetype='text/html')

@app.route('/cvparsing')
def cvparsing():
    return Response(render_template('cvparsing.html'), mimetype='text/html')

@app.route('/cvwe')
def cvwe():
    return Response(render_template('cvwe.html'), mimetype='text/html')

"""
# API ---------------------------------------------------------------------------------------------------
# """

@app.route('/api/cv/ridge/docx/extract/json', methods=['POST'])
def get_data_from_cv_json():
    # check if the post request has the file part
    if 'file' not in request.files:
        flash('No file part')
        return abort(400)

    file = request.files['file']

    # if user does not select file, browser also submit an empty part without filename
    if file.filename == '':
        flash('No selected file')
        return abort(400)

    if file and allowed_file(file.filename):
        # docx processing
        if file.filename.rsplit('.', 1)[1].lower() == 'docx':
            try:
                docx_file = secure_filename(file.filename)
                destination = "/".join([tempfile.mkdtemp(),docx_file])
                file.save(destination)
                file.close()
                if os.path.isfile(destination):
                    # return get_data_from_ridge_docx_into_json(destination)
                    return Response(get_data_from_ridge_docx_into_json(destination), mimetype='application/json')
            except Exception as e:
                logging.error(e, exc_info=True)
            return abort(500)

# --------------------------------------------------------------------------------------------------------

@app.route('/api/cv/ridge/docx/extract/xml', methods=['POST'])
def get_data_from_cv_xml():
    # check if the post request has the file part
    if 'file' not in request.files:
        flash('No file part')
        return abort(400)

    file = request.files['file']

    # if user does not select file, browser also submit an empty part without filename
    if file.filename == '':
        flash('No selected file')
        return abort(400)

    if file and allowed_file(file.filename):
        try:
            docx_file = secure_filename(file.filename)
            destination = "/".join([tempfile.mkdtemp(),docx_file])
            file.save(destination)
            file.close()
            if os.path.isfile(destination):
                # return docx_ridge_to_xml(destination, True)
                return Response(docx_ridge_to_xml(destination, True), mimetype='text/xml')
        except Exception as e:
            logging.error(e, exc_info=True)
        return abort(500)

# --------------------------------------------------------------------------------------------------------

@app.route('/api/cv/multiple/ridge/docx/extract/xml', methods=['POST'])
def get_data_from_cvs_xml():
    uploaded_files = request.files.getlist("file")
    list_xml = []
    xml_final_tree = None

    logging.debug(uploaded_files)

    for file in uploaded_files:
        if file and allowed_file(file.filename):
            # docx processing
            docx_file = secure_filename(file.filename)
            destination = "/".join([tempfile.mkdtemp(),docx_file])
            file.save(destination)
            file.close()
            if os.path.isfile(destination):
                # logging.debug(docx_ridge_to_xml(destination).getchildren())
                list_xml.append(ET.tostring(docx_ridge_to_xml(destination, False), encoding='unicode', method='xml'))
                if xml_final_tree is None:
                    xml_final_tree = docx_ridge_to_xml(destination, False)
                else:
                    for el_Node_old in xml_final_tree.iter('Node'):
                        list_xml.append(el_Node_old.get('nodeName'))
                    # ---------
                    for el_Node in docx_ridge_to_xml(destination, False).iter('Node'):
                        # xml_final_tree.find('Nodes').append(el_Node)
                        if el_Node.get('nodeName') not in list_xml:
                            xml_final_tree.find('Nodes').append(el_Node)
                    for el_Edge in docx_ridge_to_xml(destination, False).iter('Edge'):
                        xml_final_tree.find('Edges').append(el_Edge)
    return Response(ET.tostring(xml_final_tree, encoding='unicode', method='xml'), mimetype='text/xml')

@app.route('/wv/api/en/similarity', methods=['POST'])
def get_similarity():
    req_data = request.get_json()
    doc1 = NLP_NB_VECTORES(req_data['text1'])
    doc2 = NLP_NB_VECTORES(req_data['text2'])
    # doc1 = NLP_EN_VECTORES(req_data['text1'])
    # doc2 = NLP_EN_VECTORES(req_data['text2'])
    return jsonify(similarity = doc1.similarity(doc2))

# --------------------------------------------------------------------------------------------------------
@app.route('/api/bot/nb/alltermsxml', methods=['POST'])
def get_allterms():
    req_data = request.get_json()
    try:
        # spaCy doc init + default sentence normalization
        doc = NLP_NB(req_data['message'])

        # create the <allterms.xml> file structure
        # create root element <termsintext>
        root_termsintext_element = ET.Element("termsintext")
        # create element <sentences>
        sentences_element = ET.Element("sentences")
        # create element <filepath>
        # filepath_element = ET.Element("filepath")
        # filepath_element.text = file.filename
        # create element <exporterms>
        exporterms_element = ET.Element("exporterms")

        # Helper list for one-word terms
        one_word_terms_help_list = []
        # Helper list for two-word terms
        two_word_terms_help_list = []
        # Helper list for multiple-word terms (from 4-word terms)
        multiple_word_terms_help_list = []

        noun_chunks = []

        # Main text parsing cycle for sentences
        for sentence_index, sentence in enumerate(doc.sents):
            # default sentence normalization
            sentence_clean = sentence.text
            # create and append <sent>
            new_sent_element = ET.Element('sent')
            new_sent_element.text = sentence_clean #.encode('ascii', 'ignore') errors='replace'
            sentences_element.append(new_sent_element)
            # for processing specific sentence
            doc_for_chunks = NLP_NB(sentence_clean)
            # sentence NP shallow parsing cycle
            for chunk in doc_for_chunks.noun_chunks:
                doc_for_tokens = NLP_NB(chunk.text)
                '''
                # EXTRACT ONE-WORD TERMS ----------------------------------------------------------------------
                '''
                if len(doc_for_tokens) == 1:

                    if doc_for_tokens[0].pos_ in ['NOUN', 'PROPN']:

                        if doc_for_tokens[0].lemma_ in one_word_terms_help_list:
                            for term in exporterms_element.findall('term'):
                                if term.find('tname').text == doc_for_tokens[0].lemma_:
                                    new_sentpos_element = ET.Element('sentpos')
                                    new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+1)
                                    term.append(new_sentpos_element)

                        if doc_for_tokens[0].lemma_ not in one_word_terms_help_list:

                            one_word_terms_help_list.append(doc_for_tokens[0].lemma_)
                            # create and append <wcount>
                            new_wcount_element = ET.Element('wcount')
                            new_wcount_element.text = '1'
                            # create and append <ttype>
                            new_ttype_element = ET.Element('ttype')
                            new_ttype_element.text = doc_for_tokens[0].pos_
                            # create <term>
                            new_term_element = ET.Element('term')
                            # create and append <tname>
                            new_tname_element = ET.Element('tname')
                            new_tname_element.text = doc_for_tokens[0].lemma_
                            # create and append <osn>
                            new_osn_element = ET.Element('osn')
                            new_osn_element.text = NORWEGIAN_STEMMER.stem(doc_for_tokens[0].text)

                            # create and append <sentpos>
                            new_sentpos_element = ET.Element('sentpos')
                            new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+1)
                            new_term_element.append(new_sentpos_element)

                            # append to <term>
                            new_term_element.append(new_ttype_element)
                            new_term_element.append(new_tname_element)
                            new_term_element.append(new_osn_element)
                            new_term_element.append(new_wcount_element)

                            # append to <exporterms>
                            exporterms_element.append(new_term_element)

                if len(doc_for_tokens) == 2:

                    '''
                    # Extract one-word terms from 2-words statements (excluding articles DET)
                    '''
                    if doc_for_tokens[0].pos_ in ['DET', 'PUNCT']:

                        if doc_for_tokens[1].lemma_ in one_word_terms_help_list:
                            for term in exporterms_element.findall('term'):
                                if term.find('tname').text == doc_for_tokens[1].lemma_:
                                    new_sentpos_element = ET.Element('sentpos')
                                    new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+2)
                                    term.append(new_sentpos_element)

                        if doc_for_tokens[1].lemma_ not in one_word_terms_help_list:

                            one_word_terms_help_list.append(doc_for_tokens[1].lemma_)
                            # create and append <wcount>
                            new_wcount_element = ET.Element('wcount')
                            new_wcount_element.text = '1'
                            # create and append <ttype>
                            new_ttype_element = ET.Element('ttype')
                            new_ttype_element.text = doc_for_tokens[1].pos_
                            # create <term>
                            new_term_element = ET.Element('term')
                            # create and append <tname>
                            new_tname_element = ET.Element('tname')
                            new_tname_element.text = doc_for_tokens[1].lemma_
                            # create and append <osn>
                            new_osn_element = ET.Element('osn')
                            new_osn_element.text = NORWEGIAN_STEMMER.stem(doc_for_tokens[1].text)

                            # create and append <sentpos>
                            new_sentpos_element = ET.Element('sentpos')
                            new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+2)
                            new_term_element.append(new_sentpos_element)

                            # append to <term>
                            new_term_element.append(new_ttype_element)
                            new_term_element.append(new_tname_element)
                            new_term_element.append(new_osn_element)
                            new_term_element.append(new_wcount_element)

                            # append to <exporterms>
                            exporterms_element.append(new_term_element)

                    '''
                    # EXTRACT TWO-WORD TERMS ---------------------------------------------------------------
                    '''
                    if doc_for_tokens[0].pos_ not in ['DET', 'PUNCT']:

                        # print('two-word term lemma ---> ' + chunk.lemma_ +' POS[0]:'+ doc_for_tokens[0].pos_ + ' POS[0]:'+ doc_for_tokens[0].tag_ + ' HEAD[0]:' + doc_for_tokens[0].head.lower_ +' POS[1]:' + doc_for_tokens[1].pos_ + ' POS[1]:'+ doc_for_tokens[1].tag_ + ' HEAD[1]:' + doc_for_tokens[1].head.lower_)

                        # print('--------------------')

                        # If two-word term already exists in two_word_terms_help_list
                        # if chunk.lower_ in two_word_terms_help_list:
                        if chunk.lemma_ in two_word_terms_help_list:

                            # add new <sentpos> for existing two-word term
                            for term in exporterms_element.findall('term'):
                                # if term.find('tname').text == chunk.lower_:
                                if term.find('tname').text == chunk.lemma_:
                                    new_sentpos_element = ET.Element('sentpos')
                                    new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+1)
                                    term.append(new_sentpos_element)

                            # Check If root (root of Noun chunks always is a NOUN) of the two-word term
                            # already exists in one_word_terms_help_list
                            if chunk.root.lemma_ in one_word_terms_help_list:

                                sent_pos_helper = []

                                for relup_index, one_term in enumerate(exporterms_element.findall('term')):

                                    if one_term.find('tname').text == chunk.root.lemma_:

                                        for sent_pos in one_term.findall('sentpos'):
                                            sent_pos_helper.append(sent_pos.text)

                                        # create and append new <sentpos>
                                        # check if new <sentpos> already exist, if no then add new <sentpos>
                                        if chunk.root.lower_ == doc_for_tokens[0].lower_:
                                            if (str(sentence_index) + '/' + str(chunk.start+1)) not in sent_pos_helper:
                                                new_sentpos_element = ET.Element('sentpos')
                                                new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+1)
                                                one_term.append(new_sentpos_element)
                                        else:
                                            if (str(sentence_index) + '/' + str(chunk.start+2)) not in sent_pos_helper:
                                                new_sentpos_element = ET.Element('sentpos')
                                                new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+2)
                                                one_term.append(new_sentpos_element)

                            # Check If child of the root (Child not always be a NOUN, so not always be a term) of the two-word term
                            # already exists in one_word_terms_help_list
                            for t in doc_for_tokens:
                                    if t.lemma_ != chunk.root.lemma_:
                                        # if child of the root is NOUN, so it is a term
                                        if t.pos_ in ['NOUN']:
                                            if t.lemma_ in one_word_terms_help_list:

                                                sent_pos_helper = []
                                                if t.i == 0:
                                                    index_helper = chunk.start+1
                                                else:
                                                    index_helper = chunk.start+2

                                                for relup_index, one_term in enumerate(exporterms_element.findall('term')):

                                                    if one_term.find('tname').text == t.lemma_:

                                                        for sent_pos in one_term.findall('sentpos'):
                                                            sent_pos_helper.append(sent_pos.text)

                                                        if (str(sentence_index) + '/' + str(index_helper)) not in sent_pos_helper:
                                                                new_sentpos_element = ET.Element('sentpos')
                                                                new_sentpos_element.text = str(sentence_index) + '/' + str(index_helper)
                                                                one_term.append(new_sentpos_element)


                        # If two-word term not exists in two_word_terms_help_list
                        if chunk.lemma_ not in two_word_terms_help_list:

                            # update two_word_terms_help_list with the new two-word term
                            # two_word_terms_help_list.append(chunk.lower_)
                            two_word_terms_help_list.append(chunk.lemma_)

                            # create and append <wcount>
                            new_wcount_element = ET.Element('wcount')
                            new_wcount_element.text = '2'
                            # create and append <ttype>
                            new_ttype_element = ET.Element('ttype')
                            new_ttype_element.text = doc_for_tokens[0].pos_ + '_' + doc_for_tokens[1].pos_
                            # create <term>
                            new_term_element = ET.Element('term')
                            # create and append <tname>
                            new_tname_element = ET.Element('tname')
                            # new_tname_element.text = chunk.lower_
                            new_tname_element.text = chunk.lemma_
                            # create and append <osn>
                            new_osn_element = ET.Element('osn')
                            new_osn_element.text = NORWEGIAN_STEMMER.stem(doc_for_tokens[0].text)
                            new_term_element.append(new_osn_element)
                            new_osn_element = ET.Element('osn')
                            new_osn_element.text = NORWEGIAN_STEMMER.stem(doc_for_tokens[1].text)
                            new_term_element.append(new_osn_element)
                            # create and append <sentpos>
                            new_sentpos_element = ET.Element('sentpos')
                            new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+1)
                            new_term_element.append(new_sentpos_element)

                            # append to <term>
                            new_term_element.append(new_ttype_element)
                            new_term_element.append(new_tname_element)
                            new_term_element.append(new_wcount_element)

                            # append to <exporterms>
                            exporterms_element.append(new_term_element)

                            # Check If root (root of Noun chunks always is a NOUN) of the two-word term
                            # already exists in one_word_terms_help_list
                            # add relup/reldown
                            if chunk.root.lemma_ in one_word_terms_help_list:

                                sent_pos_helper = []

                                for relup_index, one_term in enumerate(exporterms_element.findall('term')):

                                    if one_term.find('tname').text == chunk.root.lemma_:

                                        for sent_pos in one_term.findall('sentpos'):
                                            sent_pos_helper.append(sent_pos.text)

                                        if chunk.root.lower_ == doc_for_tokens[0].lower_:
                                            if (str(sentence_index) + '/' + str(chunk.start+1)) not in sent_pos_helper:
                                                new_sentpos_element = ET.Element('sentpos')
                                                new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+1)
                                                one_term.append(new_sentpos_element)
                                        else:
                                            if (str(sentence_index) + '/' + str(chunk.start+2)) not in sent_pos_helper:
                                                new_sentpos_element = ET.Element('sentpos')
                                                new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+2)
                                                one_term.append(new_sentpos_element)

                                        for reldown_index, two_term in enumerate(exporterms_element.findall('term')):

                                            # if two_term.find('tname').text == chunk.lower_:
                                            if two_term.find('tname').text == chunk.lemma_:
                                                new_relup_element = ET.Element('relup')
                                                new_relup_element.text = str(relup_index)
                                                two_term.append(new_relup_element)
                                                new_reldown_element = ET.Element('reldown')
                                                new_reldown_element.text = str(reldown_index)
                                                one_term.append(new_reldown_element)

                            # Check If root NOUN not exists in one_word_terms_help_list
                            # add root NOUN to one_word_terms_help_list
                            # add relup/reldown
                            if chunk.root.lemma_ not in one_word_terms_help_list:

                                # print('root NOUN not exists in one_word_terms_help_list --->> ' + chunk.root.lemma_)
                                # print('--------------------')

                                one_word_terms_help_list.append(chunk.root.lemma_)

                                # create and append <wcount>
                                new_wcount_element = ET.Element('wcount')
                                new_wcount_element.text = '1'
                                # create and append <ttype>
                                new_ttype_element = ET.Element('ttype')
                                new_ttype_element.text = 'NOUN'
                                # create <term>
                                new_term_element = ET.Element('term')
                                # create and append <tname>
                                new_tname_element = ET.Element('tname')
                                new_tname_element.text = chunk.root.lemma_
                                # create and append <osn>
                                new_osn_element = ET.Element('osn')
                                new_osn_element.text = NORWEGIAN_STEMMER.stem(chunk.root.lower_)
                                # create and append <sentpos>
                                new_sentpos_element = ET.Element('sentpos')
                                if chunk.root.lower_ == doc_for_tokens[0].lower_:
                                    new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+1)
                                else:
                                    new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+2)
                                new_term_element.append(new_sentpos_element)
                                # append to <term>
                                new_term_element.append(new_ttype_element)
                                new_term_element.append(new_tname_element)
                                new_term_element.append(new_wcount_element)

                                # append to <exporterms>
                                exporterms_element.append(new_term_element)

                                for relup_index, one_term in enumerate(exporterms_element.findall('term')):

                                    if one_term.find('tname').text == chunk.root.lemma_:
                                        for reldown_index, two_term in enumerate(exporterms_element.findall('term')):

                                            # if two_term.find('tname').text == chunk.lower_:
                                            if two_term.find('tname').text == chunk.lemma_:
                                                new_relup_element = ET.Element('relup')
                                                new_relup_element.text = str(relup_index)
                                                two_term.append(new_relup_element)
                                                new_reldown_element = ET.Element('reldown')
                                                new_reldown_element.text = str(reldown_index)
                                                one_term.append(new_reldown_element)

                            for t in doc_for_tokens:
                                    if t.lemma_ != chunk.root.lemma_:
                                        if t.pos_ in ['NOUN']:

                                            # print('-------->>>>>>' + t.lemma_)

                                            if t.lemma_ in one_word_terms_help_list:

                                                sent_pos_helper = []
                                                if t.i == 0:
                                                    index_helper = chunk.start+1
                                                else:
                                                    index_helper = chunk.start+2


                                                for relup_index, one_term in enumerate(exporterms_element.findall('term')):

                                                    if one_term.find('tname').text == t.lemma_:
                                                        for reldown_index, two_term in enumerate(exporterms_element.findall('term')):

                                                            # if two_term.find('tname').text == chunk.lower_:
                                                            if two_term.find('tname').text == chunk.lemma_:

                                                                for sent_pos in one_term.findall('sentpos'):
                                                                    sent_pos_helper.append(sent_pos.text)

                                                                if (str(sentence_index) + '/' + str(index_helper)) not in sent_pos_helper:
                                                                    new_sentpos_element = ET.Element('sentpos')
                                                                    new_sentpos_element.text = str(sentence_index) + '/' + str(index_helper)
                                                                    one_term.append(new_sentpos_element)

                                                                new_relup_element = ET.Element('relup')
                                                                new_relup_element.text = str(relup_index)
                                                                two_term.append(new_relup_element)
                                                                new_reldown_element = ET.Element('reldown')
                                                                new_reldown_element.text = str(reldown_index)
                                                                one_term.append(new_reldown_element)

                                            if t.lemma_ not in one_word_terms_help_list:

                                                # print('if t.lemma_ not in one_word_terms_help_list ----->>>>>>' + t.lemma_)

                                                sent_pos_helper = []

                                                if t.i == 0:
                                                    index_helper = chunk.start+1
                                                else:
                                                    index_helper = chunk.start+2

                                                one_word_terms_help_list.append(t.lemma_)

                                                # create and append <wcount>
                                                new_wcount_element = ET.Element('wcount')
                                                new_wcount_element.text = '1'
                                                # create and append <ttype>
                                                new_ttype_element = ET.Element('ttype')
                                                new_ttype_element.text = 'NOUN'
                                                # create <term>
                                                new_term_element = ET.Element('term')
                                                # create and append <tname>
                                                new_tname_element = ET.Element('tname')
                                                new_tname_element.text = t.lemma_
                                                # create and append <osn>
                                                new_osn_element = ET.Element('osn')
                                                new_osn_element.text = NORWEGIAN_STEMMER.stem(t.lower_)
                                                # create and append <sentpos>
                                                new_sentpos_element = ET.Element('sentpos')
                                                new_sentpos_element.text = str(sentence_index) + '/' + str(index_helper)
                                                # append to <term>
                                                new_term_element.append(new_sentpos_element)
                                                new_term_element.append(new_ttype_element)
                                                new_term_element.append(new_tname_element)
                                                new_term_element.append(new_wcount_element)

                                                # append to <exporterms>
                                                exporterms_element.append(new_term_element)

                                                for relup_index, one_term in enumerate(exporterms_element.findall('term')):

                                                    if one_term.find('tname').text == t.lemma_:
                                                        for reldown_index, two_term in enumerate(exporterms_element.findall('term')):

                                                            # if two_term.find('tname').text == chunk.lower_:
                                                            if two_term.find('tname').text == chunk.lemma_:

                                                                for sent_pos in one_term.findall('sentpos'):
                                                                    sent_pos_helper.append(sent_pos.text)

                                                                if (str(sentence_index) + '/' + str(index_helper)) not in sent_pos_helper:
                                                                    new_sentpos_element = ET.Element('sentpos')
                                                                    new_sentpos_element.text = str(sentence_index) + '/' + str(index_helper)
                                                                    one_term.append(new_sentpos_element)

                                                                new_relup_element = ET.Element('relup')
                                                                new_relup_element.text = str(relup_index)
                                                                two_term.append(new_relup_element)
                                                                new_reldown_element = ET.Element('reldown')
                                                                new_reldown_element.text = str(reldown_index)
                                                                one_term.append(new_reldown_element)

                '''
                # EXTRACT THREE-WORD TERMS
                '''
                if len(doc_for_tokens) == 3:

                    logging.debug('three-word term lemma ---> ' + chunk.lemma_ +' POS[0]:'+ doc_for_tokens[0].pos_ + ' POS[1]:' + doc_for_tokens[1].pos_ + ' POS[2]:' + doc_for_tokens[2].pos_)
                    logging.debug('--------------------')

                if len(doc_for_tokens) > 3:

                    logging.debug('multi-word term lemma ---> ' + chunk.lemma_)
                    logging.debug('--------------------')

                    if doc_for_tokens[0].pos_ not in ['DET', 'PUNCT']:

                        # If multiple-word term already exists in multiple_word_terms_help_list
                        if chunk.lemma_ in multiple_word_terms_help_list:

                            # add new <sentpos> for existing two-word term
                            for term in exporterms_element.findall('term'):
                                if term.find('tname').text == chunk.lemma_:
                                    new_sentpos_element = ET.Element('sentpos')
                                    new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+1)
                                    term.append(new_sentpos_element)
                        
                        # If multiple-word term not exists in multiple_word_terms_help_list
                        if chunk.lemma_ not in multiple_word_terms_help_list:
                            # update  multiple_word_terms_help_list with the new multiple-word term
                            multiple_word_terms_help_list.append(chunk.lemma_)

                            # create and append <wcount>
                            new_wcount_element = ET.Element('wcount')
                            new_wcount_element.text = str(len(chunk))
                            # create and append <ttype>
                            multiple_pos_helper = []
                            for multiple_pos in doc_for_tokens:
                                multiple_pos_helper.append(multiple_pos.pos_)
                            new_ttype_element = ET.Element('ttype')
                            new_ttype_element.text = '_'.join(multiple_pos_helper)
                            # create <term>
                            new_term_element = ET.Element('term')
                            # create and append <tname>
                            new_tname_element = ET.Element('tname')
                            # new_tname_element.text = chunk.lower_
                            new_tname_element.text = chunk.lemma_
                            # create and append <osn>
                            multiple_osn_helper = []
                            for multiple_osn in doc_for_tokens:
                                new_osn_element = ET.Element('osn')
                                new_osn_element.text = NORWEGIAN_STEMMER.stem(multiple_osn.text)
                                new_term_element.append(new_osn_element)
                            # create and append <sentpos>
                            new_sentpos_element = ET.Element('sentpos')
                            new_sentpos_element.text = str(sentence_index) + '/' + str(chunk.start+1)
                            new_term_element.append(new_sentpos_element)

                            # append to <term>
                            new_term_element.append(new_ttype_element)
                            new_term_element.append(new_tname_element)
                            new_term_element.append(new_wcount_element)

                            # append to <exporterms>
                            exporterms_element.append(new_term_element)

        # create full <allterms.xml> file structure
        # root_termsintext_element.append(filepath_element)
        root_termsintext_element.append(exporterms_element)
        root_termsintext_element.append(sentences_element)

        return Response(ET.tostring(root_termsintext_element, encoding='utf8', method='xml'), mimetype='text/xml')

    except Exception as e:
        logging.error(e, exc_info=True)
        return abort(500)

@app.route('/api/bot/nb/parcexml', methods=['POST'])
def get_parcexml():
    # POS UD
    # https://universaldependencies.org/u/pos/
    if (request.args.get('pos', None) == 'ud') or (request.args.get('pos', None) == None):
        speech_dict_POS_tags = {'NOUN':'S1', 'ADJ':'S2', 'VERB': 'S4', 'INTJ':'S21', 'PUNCT':'98', 'SYM':'98', 'CONJ':'U', 'NUM':'S7', 'X':'99', 'PRON':'S11', 'ADP':'P', 'PROPN':'S22', 'ADV':'S16', 'AUX':'99', 'CCONJ':'U', 'DET':'99', 'PART':'99', 'SCONJ':'U', 'SPACE':'98'}
    # TODO Correctly relate the parts of speech with spaCy
    # POS spaCy
    if request.args.get('pos', None) == 'spacy':
        speech_dict_POS_tags = {'NOUN':'S1', 'ADJ':'S2', 'VERB': 'S4', 'INTJ':'S21', 'PUNCT':'98', 'SYM':'98', 'CONJ':'U', 'NUM':'S7', 'X':'S29', 'PRON':'S10', 'ADP':'P', 'PROPN':'S22', 'ADV':'S16', 'AUX':'AUX', 'CCONJ':'CCONJ', 'DET':'DET', 'PART':'PART', 'SCONJ':'SCONJ', 'SPACE':'SPACE'}

    req_data = request.get_json()

    try:
        doc = NLP_NB(req_data['message'])
        """
        # create the <parce.xml> file structure
        """
        # create root element <text>
        root_element = ET.Element("text")

        for sentence_index, sentence in enumerate(doc.sents):
            # XML structure creation
            new_sentence_element = ET.Element('sentence')
            # create and append <sentnumber>
            new_sentnumber_element = ET.Element('sentnumber')
            new_sentnumber_element.text = str(sentence_index+1)
            new_sentence_element.append(new_sentnumber_element)
            # create and append <sent>
            new_sent_element = ET.Element('sent')
            new_sent_element.text = sentence.text
            new_sentence_element.append(new_sent_element)

            doc_for_lemmas = NLP_NB(sentence.text)

            # create amd append <ner>, <entity>
            # NER labels description https://spacy.io/api/annotation#named-entities
            if len(doc_for_lemmas.ents) != 0:
                # create <ner>
                ner_element = ET.Element('ner')
                for ent in doc_for_lemmas.ents:
                    # create <entity>
                    new_entity_element = ET.Element('entity')
                    # create and append <entitytext>
                    new_entity_text_element = ET.Element('entitytext')
                    new_entity_text_element.text = ent.text
                    new_entity_element.append(new_entity_text_element)
                    # create and append <label>
                    new_entity_label_element = ET.Element('label')
                    new_entity_label_element.text = ent.label_
                    new_entity_element.append(new_entity_label_element)
                    # create and append <startentitypos>
                    new_start_entity_pos_character_element = ET.Element('startentityposcharacter')
                    new_start_entity_pos_token_element = ET.Element('startentitypostoken')
                    new_start_entity_pos_character_element.text = str(ent.start_char + 1)
                    new_start_entity_pos_token_element.text = str(ent.start + 1)
                    new_entity_element.append(new_start_entity_pos_character_element)
                    new_entity_element.append(new_start_entity_pos_token_element)
                    # create and append <endentitypos>
                    new_end_entity_pos_character_element = ET.Element('endentityposcharacter')
                    new_end_entity_pos_token_element = ET.Element('endentitypostoken')
                    new_end_entity_pos_character_element.text = str(ent.end_char)
                    new_end_entity_pos_token_element.text = str(ent.end)
                    new_entity_element.append(new_end_entity_pos_character_element)
                    new_entity_element.append(new_end_entity_pos_token_element)
                    # append <entity> to <ner>
                    ner_element.append(new_entity_element)
                # append <ner> to <sentence>
                new_sentence_element.append(ner_element)
            # create and append <item>, <word>, <lemma>, <number>, <pos>, <speech>
            for lemma in doc_for_lemmas:
                # create and append <item>
                new_item_element = ET.Element('item')
                # create and append <word>
                new_word_element = ET.Element('word')
                new_word_element.text = lemma.text
                new_item_element.append(new_word_element)
                # create and append <lemma>
                new_lemma_element = ET.Element('lemma')
                new_lemma_element.text = lemma.lemma_ #.encode('ascii', 'ignore')
                new_item_element.append(new_lemma_element)
                # create and append <number>
                new_number_element = ET.Element('number')
                new_number_element.text = str(lemma.i+1)
                new_item_element.append(new_number_element)
                # create and append <speech>
                new_speech_element = ET.Element('speech')
                # relate the universal dependencies parts of speech with konspekt tags
                if (request.args.get('pos', None) == 'ud') or (request.args.get('pos', None) == None):
                    # new_speech_element.text = speech_dict_POS_tags[lemma.pos_]
                    new_speech_element.text = lemma.pos_
                # relate the spaCy parts of speech with konspekt tags
                if request.args.get('pos', None) == 'spacy':
                    # new_speech_element.text = speech_dict_POS_tags[lemma.tag_]
                    new_speech_element.text = lemma.tag_
                new_item_element.append(new_speech_element)
                # create and append <pos>
                new_pos_element = ET.Element('pos')
                new_pos_element.text = str(lemma.idx+1)
                new_item_element.append(new_pos_element)

                # create and append <relate> and <rel_type>
                new_rel_type_element = ET.Element('rel_type')
                new_relate_element = ET.Element('relate')
                if lemma.dep_ == 'punct':
                    new_rel_type_element.text = 'K0'
                    new_relate_element.text = '0'
                    new_item_element.append(new_rel_type_element)
                    new_item_element.append(new_relate_element)
                else:
                    new_rel_type_element.text = lemma.dep_
                    new_item_element.append(new_rel_type_element)
                    new_relate_element.text = str(lemma.head.i+1)
                    new_item_element.append(new_relate_element)

                # create and append <group_n>
                new_group_n_element = ET.Element('group_n')
                new_group_n_element.text = '1'
                new_item_element.append(new_group_n_element)

                new_sentence_element.append(new_item_element)
            # create full <parce.xml> file structure
            root_element.append(new_sentence_element)
        return Response(ET.tostring(root_element, encoding='utf8', method='xml'), mimetype='text/xml')
    except Exception as e:
        logging.error(e, exc_info=True)
        return abort(500)
# --------------------------------------------------------------------------------------------------------

if __name__ == '__main__':
    # default port = 5000
    app.run(host = '0.0.0.0')
    # app.run(host = '127.0.0.1', port = 8000)